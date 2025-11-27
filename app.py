# ui/app.py — clean, working PoC (RAG + NLP + Upload + Generate + Insights + Business Case)

import sys, os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO

# --- DB engine (must exist before any tab code uses it)
from storage.simple_db import get_engine
engine = get_engine()

# --- NLP
from pipelines.text_nlp import analyze_transcripts

# --- Demo-mode switch (safe for free hosting) ---
import os
EE_DEPLOY_MODE = os.getenv("EE_DEPLOY_MODE", "dev")  # "demo" in the cloud

# Disable heavyweight bits when in demo:
ENABLE_VECTORS = False if EE_DEPLOY_MODE == "demo" else True

# If your storage.simple_db.get_engine uses a local SQLite file, you’re fine.
# If it expects a big local DB, fall back to in-memory in demo:
try:
    from storage.simple_db import get_engine
    engine = get_engine()
except Exception:
    from sqlalchemy import create_engine
    engine = create_engine("sqlite:///:memory:", future=True)


# ---------- Page setup ----------
st.set_page_config(page_title="CX Synthetic Insights PoC", layout="wide")
st.title("CX Synthetic Insights PoC")

# ==== VECTOR INDEX (OPTIONAL & LAZY) ====
ENABLE_VECTORS = True  # flip to False if you need to safe-boot

# Only import vector deps if enabled; never block other tabs
try:
    import storage.vector_store as vstore  # your module
    _VSTORE_IMPORT_OK = True
except Exception as e:
    _VSTORE_IMPORT_OK = False
    vstore = None
    st.sidebar.warning(f"Vector features disabled: {e}")

col = None
if ENABLE_VECTORS and _VSTORE_IMPORT_OK:
    try:
        col = vstore.get_collection()  # must handle internal init/lazy model load
        st.sidebar.success("Vector store ready")
    except Exception as e:
        st.sidebar.warning(f"Vector store unavailable: {e}")
        col = None
else:
    if not ENABLE_VECTORS:
        st.sidebar.info("Vector features are turned off for this session.")

def add_vector_index(df, dataset_id: str):
    """Index transcripts for evidence snippets (RAG). Call AFTER NLP, and only if collection exists."""
    if col is None:
        st.caption("Vector index: disabled / unavailable.")
        return
    try:
        n = vstore.upsert_embeddings(col, df, dataset_id=dataset_id)
        st.caption(f"Vector index: added ~{n} chunks")
    except Exception as e:
        st.warning(f"Vector indexing skipped: {e}")
# ==== END VECTOR INDEX ====



# ---------- Helpers ----------
def build_exec_summary_md(active_ds, mix, top_topics, pinch, high) -> str:
    lines = []
    lines.append(f"# Executive Summary — {active_ds or 'dataset'}")
    lines.append("")
    lines.append("## Highlights")
    total = int(mix["count"].sum()) if not mix.empty else 0
    neg = int(mix.loc[mix["label"] == "negative", "count"].sum()) if not mix.empty else 0
    lines.append(f"- Rows analyzed: **{total}**; Negative share: **{(neg/max(1,total)):.0%}**")
    if not top_topics.empty:
        top3 = ", ".join(top_topics.head(3)["topic"].tolist())
        lines.append(f"- Top topics: **{top3}**")
    if not pinch.empty:
        hot = pinch.sort_values(
            ["low_csat_flag", "high_wait_flag", "neg_sent_flag", "n"],
            ascending=[False, False, False, False],
        ).head(3)["stage"].tolist()
        if hot:
            lines.append(f"- Pinch points (by stage): **{', '.join(hot)}**")
    lines.append("")
    lines.append("## Pinch Points (by stage)")
    if pinch.empty:
        lines.append("_No pinch points met the thresholds yet._")
    else:
        lines.append("| stage | n | csat_mean | wait_mean | neg_share | flags |")
        lines.append("|---|---:|---:|---:|---:|---|")
        for _, r in pinch.iterrows():
            flags = ",".join([f for f, ok in {
                "low_csat": bool(r["low_csat_flag"]),
                "high_wait": bool(r["high_wait_flag"]),
                "neg_sent": bool(r["neg_sent_flag"]),
            }.items() if ok])
            lines.append(f"| {r['stage']} | {int(r['n'])} | {r['csat_mean']:.2f} | {r['wait_mean']:.1f} | {r['neg_share']:.0%} | {flags} |")
    lines.append("")
    lines.append("## High-Risk Cohort (churn_risk ≥ 2)")
    if high.empty:
        lines.append("_No high-risk interactions identified._")
    else:
        top_ht = high["topic"].value_counts().head(5).to_dict()
        top_hs = high["stage"].value_counts().head(5).to_dict()
        lines.append(f"- Size: **{len(high)}**")
        lines.append(f"- Top topics: {top_ht}")
        lines.append(f"- Top stages: {top_hs}")
    lines.append("")
    lines.append("## Suggested Actions (rule-based v1)")
    lines.append("- Pricing transparency (quotes, line items, deposits)")
    lines.append("- Booking & reminders (reduce no-shows, response time)")
    lines.append("- Upsell scripting (ethical, value-led)")
    lines.append("- Staff coaching (front desk tone, recovery)")
    lines.append("- Expectations & aftercare (outcome clarity)")
    return "\n".join(lines)

def build_docx_from_md(md_text: str) -> bytes | None:
    try:
        from docx import Document  # python-docx
    except Exception:
        return None
    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def list_datasets(engine):
    try:
        return pd.read_sql(
            """
            SELECT dataset_id,
                   COUNT(*) AS rows,
                   MIN(occurred_at) AS first_at,
                   MAX(occurred_at) AS last_at
            FROM events
            WHERE dataset_id IS NOT NULL
            GROUP BY dataset_id
            ORDER BY last_at DESC
            """,
            con=engine
        )
    except Exception:
        return pd.DataFrame(columns=["dataset_id","rows","first_at","last_at"])


# ---- Upload parser: CSV / JSON / TXT -> DataFrame ----
def parse_uploaded_file(uploaded, default_stage: str = "Discovery", default_channel: str = "phone"):
    """Returns a DataFrame with at least: customer_id, interaction_id, stage, channel, occurred_at, text"""
    name = (uploaded.name or "").lower()

    # CSV
    if name.endswith(".csv"):
        df = pd.read_csv(uploaded)

    # JSON / JSONL
    elif name.endswith(".json") or name.endswith(".jsonl"):
        try:
            df = pd.read_json(uploaded, lines=name.endswith(".jsonl"))
        except ValueError:
            try:
                uploaded.seek(0)
            except Exception:
                pass
            df = pd.read_json(uploaded)
        if isinstance(df, pd.Series):
            df = df.to_frame().T
        if "text" not in df.columns:
            for k in ["transcript", "utterance", "message", "content", "body"]:
                if k in df.columns:
                    df = df.rename(columns={k: "text"})
                    break
        if "text" not in df.columns:
            df["text"] = df.astype(str).agg(" | ".join, axis=1)

    # TXT (one line per record)
    elif name.endswith(".txt"):
        text_str = uploaded.read().decode("utf-8", errors="replace")
        lines = [ln.strip() for ln in text_str.splitlines() if ln.strip()]
        now = pd.Timestamp.utcnow().floor("s")
        df = pd.DataFrame({
            "customer_id":    [f"txt-{i:06d}" for i in range(len(lines))],
            "interaction_id": [f"txtenc-{i:06d}" for i in range(len(lines))],
            "segment":        pd.NA,
            "stage":          default_stage,
            "channel":        default_channel,
            "occurred_at":    now,
            "csat":           pd.NA,
            "wait_minutes":   pd.NA,
            "spend_cad":      pd.NA,
            "text":           lines,
        })
    else:
        return None

    # Minimal normalization
    if "text" not in df.columns: df["text"] = ""
    if "occurred_at" not in df.columns: df["occurred_at"] = pd.Timestamp.utcnow().floor("s")
    if "stage" not in df.columns: df["stage"] = default_stage
    if "channel" not in df.columns: df["channel"] = default_channel
    if "customer_id" not in df.columns: df["customer_id"] = [f"u-{i:06d}" for i in range(len(df))]
    if "interaction_id" not in df.columns: df["interaction_id"] = [f"uenc-{i:06d}" for i in range(len(df))]

    return df

# ---- Ensure required event columns exist and are well-typed ----
def ensure_event_columns(df, dataset_id: str):
    required = [
        "customer_id","interaction_id","segment","stage","channel","occurred_at",
        "csat","wait_minutes","spend_cad","text","dataset_id"
    ]
    for colname in required:
        if colname not in df.columns:
            if colname in ("customer_id","interaction_id","segment","stage","channel","text","dataset_id"):
                df[colname] = "" if colname != "segment" else pd.NA
            elif colname == "occurred_at":
                df[colname] = pd.Timestamp.utcnow().floor("s")
            else:
                df[colname] = pd.NA

    # Coerce types
    df["customer_id"]    = df["customer_id"].astype(str).fillna("")
    df["interaction_id"] = df["interaction_id"].astype(str).fillna("")
    df["segment"]        = df["segment"].astype("string")
    df["stage"]          = df["stage"].astype(str).fillna("")
    df["channel"]        = df["channel"].astype(str).fillna("")
    try:
        df["occurred_at"] = pd.to_datetime(df["occurred_at"], errors="coerce").fillna(pd.Timestamp.utcnow().floor("s"))
    except Exception:
        df["occurred_at"] = pd.Timestamp.utcnow().floor("s")
    for numcol in ("csat","wait_minutes","spend_cad"):
        df[numcol] = pd.to_numeric(df[numcol], errors="coerce")
    df["text"] = df["text"].fillna("").astype(str)
    df["dataset_id"] = dataset_id

    # Fill missing IDs if any
    if (df["interaction_id"] == "").any():
        df.loc[df["interaction_id"] == "", "interaction_id"] = [f"uenc-{i:06d}" for i in range(len(df))]
    if (df["customer_id"] == "").any():
        df.loc[df["customer_id"] == "", "customer_id"] = [f"u-{i:06d}" for i in range(len(df))]
    return df

# --- schema filter for writes
ALLOWED_COLS = [
    "customer_id","interaction_id","segment","stage","channel","occurred_at",
    "csat","wait_minutes","spend_cad","text","dataset_id",
    "sentiment_score","sentiment_label","topic","stage_guess","churn_risk",
]
# --- schema filter for writes (hardened) ---
ALLOWED_COLS = [
    "customer_id","interaction_id","segment","stage","channel","occurred_at",
    "csat","wait_minutes","spend_cad","text","dataset_id",
    "sentiment_score","sentiment_label","topic","stage_guess","churn_risk",
]

def to_sql_filtered(df: pd.DataFrame, table: str = "events"):
    # 1) keep only schema columns we care about
    keep = [c for c in ALLOWED_COLS if c in df.columns]
    df_sql = df[keep].copy()

    # 2) add any missing expected columns as NA
    for c in ALLOWED_COLS:
        if c not in df_sql.columns:
            df_sql[c] = pd.NA

    # 3) order columns
    df_sql = df_sql[ALLOWED_COLS]

    # 4) drop any accidental duplicate column names
    df_sql = df_sql.loc[:, ~pd.Index(df_sql.columns).duplicated()]

    # 5) write
    df_sql.to_sql(table, con=engine, if_exists="append", index=False)
    return df_sql

# ---------- Market TAM / Share helpers ----------
def load_yaml_safe(path: str):
    import os
    try:
        import yaml
    except Exception:
        return None
    if not os.path.exists(path):
        return None
    with open(path, "r") as f:
        try:
            return yaml.safe_load(f)
        except Exception:
            return None

def load_markets_yaml(path: str = "config/markets.yml") -> dict:
    data = load_yaml_safe(path) or {}
    return data.get("markets", {})

def load_market_levers_yaml(path: str = "config/market_levers.yml") -> dict:
    data = load_yaml_safe(path) or {}
    return data.get("levers", {})

def _pp_from_range(rng: list | tuple, intensity_0_1: float) -> float:
    """Linear interpolate a [min,max] range by intensity (0..1)."""
    try:
        lo, hi = float(rng[0]), float(rng[1])
        t = max(0.0, min(1.0, float(intensity_0_1)))
        return lo + (hi - lo) * t
    except Exception:
        return 0.0

def compute_market_scenarios(
    markets_cfg: dict,
    target_shares_pct: list[int] = [1, 3],
    lever_uplift_pp: float = 0.0
) -> pd.DataFrame:
    """
    Build a tidy DF with baseline vs 2 targets and an optional 'lever-based' target:
    projected_share = baseline_share_pct + lever_uplift_pp
    """
    rows = []
    for key, m in markets_cfg.items():
        label = m.get("label", key)
        currency = m.get("currency", "USD")
        tam = float(m.get("wearable_tam", 0.0))
        reachable_pct = float(m.get("reachable_pct", 1.0))
        baseline_share_pct = float(m.get("baseline_share_pct", 0.0))
        gm_pct = float(m.get("gross_margin_pct", 0.0))

        reachable_rev = tam * reachable_pct
        baseline_rev = reachable_rev * (baseline_share_pct / 100.0)
        baseline_gp = baseline_rev * gm_pct

        # standard targets (e.g., 1% & 3% of reachable TAM)
        for t in sorted({int(v) for v in (target_shares_pct or []) if v is not None}):
            target_share_pct = max(0.0, float(t))
            target_rev = reachable_rev * (target_share_pct / 100.0)
            target_gp = target_rev * gm_pct
            rows.append({
                "region_key": key,
                "Region": label,
                "Currency": currency,
                "Scenario": f"{int(target_share_pct)}% Share",
                "Baseline_share_pct": baseline_share_pct,
                "Projected_share_pct": target_share_pct,
                "Reachable_TAM": reachable_rev,
                "Baseline_revenue": baseline_rev,
                "Target_revenue": target_rev,
                "Incremental_revenue": target_rev - baseline_rev,
                "Baseline_gross_profit": baseline_gp,
                "Target_gross_profit": target_gp,
                "Incremental_gross_profit": target_gp - baseline_gp,
                "Gross_margin_pct": gm_pct,
            })

        # lever-only target (baseline + uplift)
        if lever_uplift_pp and lever_uplift_pp > 0:
            proj_share = max(0.0, baseline_share_pct + float(lever_uplift_pp))
            target_rev = reachable_rev * (proj_share / 100.0)
            target_gp = target_rev * gm_pct
            rows.append({
                "region_key": key,
                "Region": label,
                "Currency": currency,
                "Scenario": f"Levers (+{lever_uplift_pp:.1f}pp)",
                "Baseline_share_pct": baseline_share_pct,
                "Projected_share_pct": proj_share,
                "Reachable_TAM": reachable_rev,
                "Baseline_revenue": baseline_rev,
                "Target_revenue": target_rev,
                "Incremental_revenue": target_rev - baseline_rev,
                "Baseline_gross_profit": baseline_gp,
                "Target_gross_profit": target_gp,
                "Incremental_gross_profit": target_gp - baseline_gp,
                "Gross_margin_pct": gm_pct,
            })

    df = pd.DataFrame(rows)
    if not df.empty:
        df = df.sort_values(["Region", "Scenario"])
    return df

# ---------- Market levers & scenarios helpers ----------
def load_market_levers_yaml(path: str = "config/market_levers.yml"):
    import os
    try:
        import yaml
    except Exception:
        return {}
    if not os.path.exists(path):
        return {}
    with open(path, "r") as f:
        data = yaml.safe_load(f) or {}
    return data.get("levers", {})

def load_market_scenarios_yaml(path: str = "config/market_scenarios.yml"):
    import os
    try:
        import yaml
    except Exception:
        return {}
    if not os.path.exists(path):
        return {}
    with open(path, "r") as f:
        data = yaml.safe_load(f) or {}
    return data.get("scenarios", {})

def _lever_effect_pp(levers_cfg: dict, lever_key: str, mode: str, intensity: float) -> float:
    """
    Return share uplift (percentage points) for a single lever given a 'mode'
    (Conservative / Typical / Aggressive) and intensity 0..1.
    """
    meta = levers_cfg.get(lever_key, {})
    rng = meta.get("effect_on_share_pp", [0.0, 0.0])
    lo = float(rng[0]) if len(rng) > 0 else 0.0
    hi = float(rng[1]) if len(rng) > 1 else lo
    if mode == "Conservative":
        base = lo
    elif mode == "Aggressive":
        base = hi
    else:
        base = (lo + hi) / 2.0
    return max(0.0, min(1.0, float(intensity))) * base


# ---------- Tabs ----------
tab_gen, tab_journey, tab_insights, tab_business, tab_market, tab_playbook = st.tabs(
    ["Ingest / Generate", "Journey & Explore", "Insights & Actions",
     "Business Case", "Market Expansion", "Playbook"]
)

# =========================================================
# Tab 1: Ingest / Generate
# =========================================================
with tab_gen:
    st.subheader("Generate synthetic data (Medispa) — segments, spend (CAD) & repeat")

    # dataset id/session
    if "active_dataset_id" not in st.session_state:
        st.session_state["active_dataset_id"] = f"trackA_{pd.Timestamp.utcnow().strftime('%Y%m%d-%H%M%S')}"
    dataset_name = st.text_input("Dataset name", st.session_state["active_dataset_id"])

    # Import generator with inline debug
    try:
        from generators.track_a import TrackAParams, generate as gen_a
        import inspect
        st.caption(f"Generator path: {inspect.getsourcefile(gen_a)}")
        st.caption(f"TrackAParams fields: {list(getattr(TrackAParams, '__annotations__', {}).keys())}")
        gen_import_ok = True
    except Exception as e:
        gen_import_ok = False
        st.error(f"Generator import error: {e}\n\nMake sure generators/track_a.py exists and defines TrackAParams.")

    if gen_import_ok:
        # basic controls
        seed = st.number_input("Random seed (reproducible)", value=42, step=1)
        n_customers = st.number_input("Customers to simulate", value=500, min_value=50, step=50)
        median_wait = st.number_input("Median wait (minutes)", value=12.0, min_value=0.0, step=0.5)
        sigma       = st.number_input("Wait time spread (σ)", value=6.0, min_value=0.0, step=0.5)
        csat_base   = st.number_input("Baseline CSAT (1–5)", value=4.0, min_value=1.0, max_value=5.0, step=0.1)

        # Generate
        if st.button("Generate Track A dataset"):
            import inspect
            candidate_kwargs = dict(
                n_customers=int(n_customers),
                seed=int(seed),
                median_wait_min=float(median_wait),
                wait_sigma=float(sigma),
                csat_base=float(csat_base),
            )
            allowed = set(inspect.signature(TrackAParams).parameters.keys())
            safe_kwargs = {k: v for k, v in candidate_kwargs.items() if k in allowed and candidate_kwargs[k] is not None}
            params = TrackAParams(**safe_kwargs)

            df = gen_a(params)
            df["dataset_id"] = dataset_name
            st.session_state["active_dataset_id"] = dataset_name

            if "text" not in df.columns: df["text"] = ""
            df["text"] = df["text"].fillna("")

            # NLP + vectors
            df = analyze_transcripts(df)
            add_vector_index(df, dataset_name)

            # Persist (filtered)
            df_sql = to_sql_filtered(df, table="events")

            st.success(f"Generated {len(df_sql)} rows into dataset '{dataset_name}'.")
            st.write(df_sql.head())

    st.divider()
    st.subheader("Upload CSV / JSON / TXT")

    upload_mode_label = st.radio("Upload mode", ["Append", "Replace"], horizontal=True)

    with st.expander("Defaults for plain-text uploads (.txt) and fallback for JSON/CSV without fields"):
        default_stage = st.selectbox(
            "Default stage",
            ["Discovery","Consult_Booking","Consult_Followup","Treatment",
             "Support_QA","Ongoing_Comms","Repeat_Booking"],
            index=1,
        )
        default_channel = st.selectbox("Default channel", ["in_person","phone","web","referral"], index=1)

    uploaded = st.file_uploader("Upload CSV / JSON / TXT", type=["csv","json","txt"])

    if uploaded and st.button("Upload file"):
        try:
            # 1) parse
            df_up = parse_uploaded_file(uploaded, default_stage=default_stage, default_channel=default_channel)
            if df_up is None or df_up.empty:
                st.error("Parsed upload is empty. Please check the file content.")
            else:
                # 2) choose target dataset
                target_ds = st.session_state.get("active_dataset_id", "dataset_untitled")

                # 3) ensure schema + tag
                df_up = ensure_event_columns(df_up, dataset_id=target_ds)

                # 4) NLP
                df_up = analyze_transcripts(df_up)

                # 5) (optional) map external fields -> our schema if present
                rename_map = {"id": "interaction_id", "participant_id": "customer_id", "date": "occurred_at"}
                df_up = df_up.rename(columns={k: v for k, v in rename_map.items() if k in df_up.columns})

                # 6) vectors (full df for metadata richness)
                add_vector_index(df_up, target_ds)

                # 7) Replace mode delete
                if upload_mode_label == "Replace":
                    import sqlalchemy as sa
                    with engine.begin() as conn:
                        conn.execute(sa.text("DELETE FROM events WHERE dataset_id = :d"), {"d": target_ds})

                # 8) persist (filtered)
                df_sql = to_sql_filtered(df_up, table="events")

                st.session_state["active_dataset_id"] = target_ds
                st.success(f"Ingested {len(df_sql)} rows into dataset '{target_ds}' (mode={upload_mode_label}).")
                st.write(df_sql.head())
        except Exception as e:
            st.error(f"Upload failed: {e}")


# =========================================================
# Tab 2: Journey & Explore
# =========================================================
with tab_journey:
    st.subheader("Explore journey events")

    # --- Dataset picker (Active Dataset) ---
    try:
        all_ids = pd.read_sql(
            "SELECT DISTINCT dataset_id FROM events WHERE dataset_id IS NOT NULL ORDER BY dataset_id DESC",
            con=engine
        )
        options = all_ids["dataset_id"].dropna().tolist()
    except Exception:
        options = []

    if options:
        default_idx = 0
        if st.session_state.get("active_dataset_id") in options:
            default_idx = options.index(st.session_state["active_dataset_id"])
        active_ds = st.selectbox("Active dataset", options, index=default_idx)
        st.session_state["active_dataset_id"] = active_ds
        st.caption(f"Analyzing dataset: **{active_ds}**")
    else:
        st.info("No datasets found. Generate or upload on the first tab.")
        active_ds = None

    if not active_ds:
        st.stop()

    # --- Show datasets & quick switch to newest ---
    try:
        ds_table = pd.read_sql(
            """
            SELECT dataset_id,
                   COUNT(*) AS rows,
                   MIN(occurred_at) AS first_at,
                   MAX(occurred_at) AS last_at
            FROM events
            WHERE dataset_id IS NOT NULL
            GROUP BY dataset_id
            ORDER BY last_at DESC
            """,
            con=engine
        )
    except Exception:
        ds_table = pd.DataFrame(columns=["dataset_id","rows","first_at","last_at"])

    if not ds_table.empty:
        st.markdown("**Datasets & Row Counts**")
        st.dataframe(ds_table, use_container_width=True)

        freshest = ds_table.iloc[0]["dataset_id"]
        if st.button(f"Switch to newest: {freshest}", key="switch_freshest"):
            st.session_state["active_dataset_id"] = freshest
            try:
                st.rerun()
            except Exception:
                try:
                    st.experimental_rerun()
                except Exception:
                    st.warning("Please click the ↻ Rerun button at the top.")

    # --- Load ONLY the active dataset ---
    try:
        df = pd.read_sql(
            "SELECT * FROM events WHERE dataset_id = :d ORDER BY occurred_at DESC LIMIT 20000",
            con=engine,
            params={"d": active_ds}
        )
        # 🔎 quick debug
        st.caption(f"Rows loaded for {active_ds}: {len(df)}")
        with st.expander("Peek first 5 rows (debug)", expanded=False):
            st.dataframe(df.head(), use_container_width=True)
    except Exception as e:
        st.info("There was an issue loading data from SQLite.")
        st.caption(f"(debug) Journey load error: {e}")
        df = pd.DataFrame()

    # --- KPIs & visuals ---
    if df.empty:
        st.info("No rows in this dataset yet. Generate or upload on the first tab.")
    else:
        # Ensure numeric types
        for coln in ["csat", "wait_minutes", "spend_cad"]:
            if coln in df.columns:
                df[coln] = pd.to_numeric(df[coln], errors="coerce")

        spend_col = "spend_cad" if "spend_cad" in df.columns else None

        # KPIs
        k1, k2, k3, k4 = st.columns(4)
        with k1:
            st.metric("Events", f"{len(df):,}")
        with k2:
            st.metric("Customers", f"{df['customer_id'].nunique():,}" if "customer_id" in df.columns else "—")
        with k3:
            st.metric("Avg CSAT", f"{df['csat'].mean():.2f}" if "csat" in df.columns and df['csat'].notna().any() else "—")
        with k4:
            if spend_col:
                st.metric("Total Spend (CAD)", f"{df[spend_col].fillna(0).sum():,.0f}")
            else:
                st.metric("Total Spend", "—")

        # Stage summary
        st.markdown("### Stage summary")
        agg = df.groupby("stage", dropna=True).agg(
            events=("interaction_id","count"),
            avg_wait=("wait_minutes","mean") if "wait_minutes" in df.columns else ("interaction_id","count"),
            avg_csat=("csat","mean") if "csat" in df.columns else ("interaction_id","count"),
            total_spend=(spend_col,"sum") if spend_col else ("interaction_id","count"),
        ).reset_index().sort_values("events", ascending=False)
        st.dataframe(agg, use_container_width=True)

        left, right = st.columns(2)
        with left:
            st.markdown("**Events by stage**")
            st.bar_chart(agg.set_index("stage")["events"])
        with right:
            if spend_col:
                st.markdown("**Spend by stage (CAD)**")
                st.bar_chart(agg.set_index("stage")["total_spend"])

        # Segment insights (if present)
        if "segment" in df.columns:
            st.markdown("### Segment insights")
            seg_df = df.copy()
            seg_df["is_repeat"] = (seg_df["stage"] == "Repeat_Booking").astype(int)

            mix = seg_df.groupby("segment", dropna=True).agg(
                customers=("customer_id","nunique") if "customer_id" in seg_df.columns else ("interaction_id","count"),
                events=("interaction_id","count"),
                repeat_events=("is_repeat","sum"),
                spend_cad=(spend_col,"sum") if spend_col else ("interaction_id","count"),
                avg_csat=("csat","mean") if "csat" in seg_df.columns else ("interaction_id","count"),
            ).reset_index().sort_values("customers", ascending=False)
            st.dataframe(mix, use_container_width=True)

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Customers by segment**")
                st.bar_chart(mix.set_index("segment")["customers"])
            with c2:
                st.markdown("**Repeat events by segment**")
                st.bar_chart(mix.set_index("segment")["repeat_events"])

        # Drilldown
        st.markdown("### Drilldown")
        st.caption("Filter quickly by stage/segment; click column headers to sort.")
        cstage, cseg = st.columns(2)
        stage_sel = cstage.multiselect("Stage", sorted(df["stage"].dropna().unique().tolist()) if "stage" in df.columns else [])
        seg_sel   = cseg.multiselect("Segment", sorted(df["segment"].dropna().unique().tolist()) if "segment" in df.columns else [])
        ddf = df.copy()
        if stage_sel:
            ddf = ddf[ddf["stage"].isin(stage_sel)]
        if "segment" in ddf.columns and seg_sel:
            ddf = ddf[ddf["segment"].isin(seg_sel)]
        st.dataframe(ddf.sort_values("occurred_at", ascending=False).head(1000), use_container_width=True)




# =========================================================
# Tab 3: Insights & Actions
# =========================================================
with tab_insights:
    st.subheader("Insights v1")

    try:
        df_all = pd.read_sql("SELECT * FROM events", con=engine)
        if df_all.empty:
            st.info("No data yet. Generate or upload a dataset first.")
        else:
            active_ds = st.session_state.get("active_dataset_id")
            df = df_all[df_all["dataset_id"] == active_ds].copy() if active_ds else df_all.copy()

            st.caption(f"Active dataset: {active_ds} — {len(df)} rows")

            # RAG backfill (one-time per dataset)
            if col is not None:
                if st.button("Index current dataset for evidence"):
                    added = vstore.upsert_embeddings(col, df, dataset_id=active_ds or "default")
                    st.success(f"Indexed ~{added} chunks for dataset '{active_ds}'.")

            # Sentiment mix
            st.markdown("**Sentiment mix**")
            mix = df["sentiment_label"].value_counts().rename_axis("label").reset_index(name="count")
            st.dataframe(mix, use_container_width=True)
            st.bar_chart(mix.set_index("label"))

            # Top topics
            st.markdown("**Top topics**")
            top_topics = df["topic"].value_counts().rename_axis("topic").reset_index(name="count")
            st.dataframe(top_topics, use_container_width=True)
            st.bar_chart(top_topics.set_index("topic"))

            # Pinch points (by stage)
            st.markdown("**Pinch points (by stage)**")
            THRESHOLDS = {"csat_low": 3.4, "wait_high": 10, "neg_share": 0.30, "min_n": 5}
            grp = df.groupby("stage", dropna=False).agg(
                n=("interaction_id", "count"),
                csat_mean=("csat", "mean"),
                wait_mean=("wait_minutes", "mean"),
                neg_share=("sentiment_label", lambda s: (s == "negative").mean()),
            ).reset_index()
            grp["low_csat_flag"]  = (grp["n"] >= THRESHOLDS["min_n"]) & (grp["csat_mean"] < THRESHOLDS["csat_low"])
            grp["high_wait_flag"] = (grp["n"] >= THRESHOLDS["min_n"]) & (grp["wait_mean"] > THRESHOLDS["wait_high"])
            grp["neg_sent_flag"]  = (grp["n"] >= THRESHOLDS["min_n"]) & (grp["neg_share"] > THRESHOLDS["neg_share"])
            pinch = grp[(grp["low_csat_flag"]) | (grp["high_wait_flag"]) | (grp["neg_sent_flag"])].copy()
            st.dataframe(
                pinch.sort_values(["low_csat_flag","high_wait_flag","neg_sent_flag","n"], ascending=[False,False,False,False]),
                use_container_width=True,
            )

            # High-risk cohort
            st.markdown("**High-risk cohort (churn_risk ≥ 2)**")
            high = df[df["churn_risk"] >= 2].copy()
            st.write({
                "cohort_size": int(len(high)),
                "share_of_dataset": round(len(high) / max(1, len(df)), 3),
                "top_topics": high["topic"].value_counts().head(5).to_dict(),
                "top_stages": high["stage"].value_counts().head(5).to_dict(),
            })
            st.dataframe(
                high[["customer_id","interaction_id","stage","topic","sentiment_label","csat","wait_minutes","text"]].head(15),
                use_container_width=True,
            )

            # CSV downloads
            def _csv_bytes(xdf): return xdf.to_csv(index=False).encode("utf-8")
            c1, c2, c3, c4 = st.columns(4)
            with c1: st.download_button("Download Sentiment CSV", _csv_bytes(mix), "sentiment_mix.csv", "text/csv")
            with c2: st.download_button("Download Topics CSV", _csv_bytes(top_topics), "top_topics.csv", "text/csv")
            with c3: st.download_button("Download Pinch Points CSV", _csv_bytes(pinch), "pinch_points.csv", "text/csv")
            with c4:
                st.download_button(
                    "Download High-Risk CSV",
                    _csv_bytes(high[["customer_id","interaction_id","stage","topic","sentiment_label","csat","wait_minutes","text"]]),
                    "high_risk_cohort.csv", "text/csv",
                )

            # Executive one-pager
            st.markdown("### Export Executive Summary")
            md_text = build_exec_summary_md(active_ds, mix, top_topics, pinch, high)
            colA, colB, colC = st.columns(3)
            with colA:
                st.download_button("Download .md", md_text.encode("utf-8"),
                                   file_name=f"exec_summary_{active_ds or 'dataset'}.md", mime="text/markdown")
            with colB:
                st.download_button("Download .txt", md_text.encode("utf-8"),
                                   file_name=f"exec_summary_{active_ds or 'dataset'}.txt", mime="text/plain")
            with colC:
                docx_bytes = build_docx_from_md(md_text)
                if docx_bytes:
                    st.download_button("Download .docx", docx_bytes,
                        file_name=f"exec_summary_{active_ds or 'dataset'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.caption("Install `python-docx` for .docx export (optional).")

            # Evidence snippets (RAG)
            show_ev = (col is not None) and st.checkbox("Show evidence snippets", value=True)
            def _evidence_for(topic: str | None, stage: str | None, n: int = 5):
                if col is None: return []
                q = " ".join([
                    (topic or "").replace("_"," "),
                    (stage or "").replace("_"," "),
                    "customer complaint feedback comment issue"
                ]).strip()
                hits = vstore.similarity_search(col, query=q, dataset_id=active_ds, k=25)
                # soft filter
                if topic: hits = [h for h in hits if topic.lower() in (h.get("metadata",{}).get("topic","").lower())]
                if stage: hits = [h for h in hits if stage.lower() in (h.get("metadata",{}).get("stage","").lower())]
                return hits[:n]
            if show_ev:
                st.markdown("### Evidence: Pinch points")
                if pinch.empty:
                    st.caption("No pinch points flagged yet.")
                else:
                    for _, row in pinch.sort_values(
                        ["low_csat_flag","high_wait_flag","neg_sent_flag","n"],
                        ascending=[False,False,False,False]
                    ).head(3).iterrows():
                        stage_i = str(row["stage"])
                        top_topic = (df[df["stage"]==stage_i]["topic"].value_counts().index[0]
                                     if (df["stage"]==stage_i).any() else None)
                        with st.expander(f"{stage_i} — evidence (topic: {top_topic or 'mixed'})", expanded=False):
                            for hit in _evidence_for(top_topic, stage_i, n=5):
                                meta = hit.get("metadata", {})
                                st.write(f"• {hit['document']}")
                                st.caption(f"{meta.get('stage','?')} · {meta.get('topic','?')} · {meta.get('sentiment_label','?')} · {meta.get('source_pointer','')}")
                st.markdown("### Evidence: High-risk cohort")
                if high.empty:
                    st.caption("No high-risk items.")
                else:
                    with st.expander("Top snippets (high risk)", expanded=False):
                        for hit in _evidence_for(None, None, n=5):
                            meta = hit.get("metadata", {})
                            st.write(f"• {hit['document']}")
                            st.caption(f"{meta.get('stage','?')} · {meta.get('topic','?')} · {meta.get('sentiment_label','?')} · {meta.get('source_pointer','')}")
    except Exception as e:
        st.error(f"Insights failed to load: {e}")


# =========================================================
# Tab 4: Business Case (Guided, Staff Optimization ROI)
# =========================================================
with tab_business:
    st.subheader("Business Case — Staff Optimization (no new hires)")

    # ------------------ STEP 0: Presets ------------------
    PRESETS = {
        "Medispa — Mark (Small Clinic)": {
            # Anchors
            "annual_spend_per_customer": 1200.0,   # $1.5M / ~1250 ≈ $1200
            "avg_ticket": 300.0,                   # AOV per visit
            "customers_per_year": 1250,            # from baseline
            "avg_clv": 1600.0,                     # simple revenue-based CLV anchor

            # Baseline funnel
            "conv_pct": 70,                        # consult → treatment (Mark)
            "no_show_pct": 10,                     # starter if unknown

            # Time & rework
            "minutes_per_visit": 35,
            "rework_pct": 10,
            "rework_minutes": 20,

            # Lifetime effect defaults
            "use_lifetime": True,
            "base_csat": 4.0,
            "csat_delta": 0.0,                     # leave 0 if unknown
            "elasticity": 0.12,

            # Capacity / new-visit conversion
            "fill_rate_pct": 60,
            "minutes_per_new_visit": 35,
        },

        "Medispa – Value (Pack A)": {
            "annual_spend_per_customer": 1200.0,
            "avg_ticket": 300.0,
            "customers_per_year": 1250,
            "conv_pct": 55,
            "no_show_pct": 12,
            "minutes_per_visit": 35,
            "rework_pct": 10,
            "rework_minutes": 20,
            "use_lifetime": True,
            "base_csat": 4.0,
            "csat_delta": 0.2,
            "elasticity": 0.12,
            "fill_rate_pct": 60,
            "minutes_per_new_visit": 35,
            "avg_clv": 1600.0,
        },

        "Medispa – Premium": {
            "annual_spend_per_customer": 1800.0,
            "avg_ticket": 450.0,
            "customers_per_year": 1000,
            "conv_pct": 60,
            "no_show_pct": 8,
            "minutes_per_visit": 40,
            "rework_pct": 8,
            "rework_minutes": 25,
            "use_lifetime": True,
            "base_csat": 4.2,
            "csat_delta": 0.25,
            "elasticity": 0.12,
            "fill_rate_pct": 65,
            "minutes_per_new_visit": 40,
            "avg_clv": 2200.0,
        },
    }

    def _apply_preset_to_state(preset_name: str):
        p = PRESETS[preset_name]
        st.session_state.setdefault("roi_state", {})
        st.session_state["roi_state"].update(p)
        st.session_state["roi_state"]["preset"] = preset_name

    # Auto-apply default once
    if "roi_state" not in st.session_state:
        _apply_preset_to_state("Medispa — Mark (Small Clinic)")

    state = st.session_state["roi_state"]

    c0a, c0b = st.columns([2, 1])
    with c0a:
        preset_choice = st.selectbox(
            "Step 0 — Choose a preset",
            list(PRESETS.keys()),
            index=list(PRESETS.keys()).index(state.get("preset", "Medispa — Mark (Small Clinic)")),
        )
    with c0b:
        if st.button("Apply preset"):
            _apply_preset_to_state(preset_choice)
            st.rerun()


    state = st.session_state["roi_state"]

    # --- Tab 4: widget key bootstrap (must run before any widgets use these keys)
if "roi_defaults_bootstrapped" not in st.session_state:
    # Core-8 lever keys
    st.session_state.setdefault("aov_mode_key", "Model via add-ons in Lever 4")
    st.session_state.setdefault("aov_delta_pct", 0)
    st.session_state.setdefault("visit_freq_delta_pct", 0)
    st.session_state.setdefault("d_conv_pp", 0)
    st.session_state.setdefault("d_no_show_pp", 0)
    st.session_state.setdefault("d_handle_time_min", 0)
    st.session_state.setdefault("d_rework_pp", 0)
    st.session_state.setdefault("fill_rate_pct", state.get("fill_rate_pct", 60))
    st.session_state.setdefault("attach_rate_delta_pct", 0)
    st.session_state.setdefault("avg_addon_value", 120.0)

    # Costs
    st.session_state.setdefault("tool_cost_mo", 0.0)
    st.session_state.setdefault("project_cost_one_time", 0.0)

    # Misc / notes (used in the playbook tab)
    st.session_state.setdefault("notes_keys", {})

    st.session_state["roi_defaults_bootstrapped"] = True

# If a playbook tab queued updates for ROI, apply them NOW (before widgets render)
pending = st.session_state.pop("pending_roi_updates", None)
if pending:
    for k, v in pending.items():
        st.session_state[k] = v


    # ------------------ STEP 1: Economic anchors ------------------
    st.markdown("### Step 1 — Economic anchors")
    st.caption("Pick values; we derive visits/customer/year and visits/month. You can override visits/month later.")

    a1, a2, a3, a4 = st.columns([1,1,1,1])
    with a1:
        annual_spend_per_customer = st.number_input(
            "Avg annual spend / customer (CAD)",
            value=float(state.get("annual_spend_per_customer", 1200.0)),
            min_value=0.0, step=50.0
        )
    with a2:
        avg_ticket_anchor = st.number_input(
            "Avg ticket (per visit, CAD)",
            value=float(state.get("avg_ticket", 300.0)),
            min_value=0.0, step=10.0
        )
    with a3:
        customers_per_year = st.number_input(
            "Active customers per year",
            value=int(state.get("customers_per_year", 1250)),
            min_value=0, step=50
        )
    with a4:
        avg_clv_anchor = st.number_input(
            "Avg CLV (CAD) (optional; revenue-based for now)",
            value=float(state.get("avg_clv", 1600.0)),
            min_value=0.0, step=50.0
        )

    avg_clv_base = float(state.get("avg_clv", avg_clv_anchor))


    visits_per_customer_per_year = (annual_spend_per_customer / avg_ticket_anchor) if avg_ticket_anchor > 0 else 0.0
    visits_per_year = customers_per_year * visits_per_customer_per_year
    visits_derived = int(round(visits_per_year / 12.0))
    st.caption(
        f"Derived: visits/customer/year = {visits_per_customer_per_year:.2f} → visits/year = {visits_per_year:,.0f} → visits/month ≈ {visits_derived:,}"
    )

    # ------------------ STEP 2: Baseline & time ------------------
    st.markdown("### Step 2 — Baseline & time")
    c1, c2 = st.columns(2)
    with c1:
        visits = st.number_input("Scheduled visits per month (derived; adjust if needed)",
                                 value=int(state.get("visits", visits_derived)),
                                 min_value=0, step=10)
        avg_ticket = st.number_input("Avg ticket used in ROI (CAD)",
                                     value=float(state.get("avg_ticket", avg_ticket_anchor)), min_value=0.0, step=10.0)

        consult_to_treatment_conv = st.slider("Consult → Treatment conversion (%)", 0, 100,
                                              int(state.get("conv_pct", 55)), 1)
        no_show_rate = st.slider("No-show rate (%)", 0, 100,
                                 int(state.get("no_show_pct", 12)), 1)

    with c2:
        minutes_per_visit = st.number_input("Avg staff minutes per completed visit",
                                            value=int(state.get("minutes_per_visit", 35)), min_value=1, step=1)
        avg_rework_rate   = st.slider("Rework rate (% needing correction/follow-up)", 0, 100,
                                      int(state.get("rework_pct", 10)), 1)
        avg_rework_minutes = st.number_input("Avg staff minutes per rework event",
                                             value=int(state.get("rework_minutes", 20)), min_value=0, step=1)

        use_lifetime = st.checkbox("Include lifetime effects (CSAT → retention)",
                                   value=bool(state.get("use_lifetime", True)))
        base_csat    = st.slider("Baseline CSAT", 1.0, 5.0, float(state.get("base_csat", 4.0)), 0.1)
        csat_delta   = st.number_input("Expected CSAT delta (Δ from baseline)",
                                       value=float(state.get("csat_delta", 0.2)), step=0.05)
        elasticity   = st.number_input("CSAT → retention elasticity (per +1.0)",
                                       value=float(state.get("elasticity", 0.12)), min_value=0.0, max_value=1.0, step=0.01)

    # ---- Solutions catalog (YAML) helper: safe + self-contained ----
    def load_solutions_yaml(path: str = "config/solutions.yml"):
        import os
        try:
            import yaml  # PyYAML
        except Exception:
            yaml = None

        if yaml is None or not os.path.exists(path):
            return []  # no library or no file → no presets

        with open(path, "r") as f:
            data = (yaml.safe_load(f) or [])

        # Core-8 fields we accept from YAML
        allowed = {
            "aov_delta_pct",            # Lever 1 — direct AOV %
            "visit_freq_delta_pct",     # Lever 2 — scheduled visits %
            "d_conv_pp",                # Lever 3
            "attach_rate_delta_pp",     # Lever 4
            "avg_addon_value",          # Lever 4
            "d_no_show_pp",             # Lever 5
            "csat_delta",               # Lever 6
            "d_handle_time_min",        # Lever 7
            "d_rework_pp",              # Lever 7
            "fill_rate_pct",            # Lever 7
        }

        cleaned = []
        for s in data:
            s = dict(s)
            s["levers"] = {k: v for k, v in (s.get("levers") or {}).items() if k in allowed}
            cleaned.append(s)
        return cleaned


    def _aggregate_levers(selected: list, base_fill_rate: int) -> dict:
        """Sum additive deltas, max() the fill rate; defaults if missing."""
        agg = {
            "aov_delta_pct": 0,
            "visit_freq_delta_pct": 0,
            "d_conv_pp": 0,
            "attach_rate_delta_pp": 0,
            "avg_addon_value": 0.0,
            "d_no_show_pp": 0,
            "csat_delta": 0.0,
            "d_handle_time_min": 0,
            "d_rework_pp": 0,
            "fill_rate_pct": int(base_fill_rate),
        }
        for s in selected or []:
            L = (s.get("levers") or {})
            agg["aov_delta_pct"]          += int(L.get("aov_delta_pct", 0))
            agg["visit_freq_delta_pct"]   += int(L.get("visit_freq_delta_pct", 0))
            agg["d_conv_pp"]              += int(L.get("d_conv_pp", 0))
            agg["attach_rate_delta_pp"]   += int(L.get("attach_rate_delta_pp", 0))
            agg["avg_addon_value"]        += float(L.get("avg_addon_value", 0.0))
            agg["d_no_show_pp"]           += int(L.get("d_no_show_pp", 0))
            agg["csat_delta"]             += float(L.get("csat_delta", 0.0))
            agg["d_handle_time_min"]      += int(L.get("d_handle_time_min", 0))
            agg["d_rework_pp"]            += int(L.get("d_rework_pp", 0))
            agg["fill_rate_pct"]           = max(agg["fill_rate_pct"], int(L.get("fill_rate_pct", agg["fill_rate_pct"])))
        return agg

    # ------------------ STEP 3: Core-8 Improvement Levers (numbered) ------------------
    st.markdown("### Step 3 — Core-8 Improvement Levers (set expected changes)")

    # Optional: apply solution presets from YAML to prefill the sliders
    solutions_catalog = load_solutions_yaml("config/solutions.yml")
    if solutions_catalog:
        with st.expander("Apply solution presets from catalog (optional)", expanded=False):
            names = [s["name"] for s in solutions_catalog]
            sel_names = st.multiselect("Choose solution(s)", names, key="solutions_sel_v2")
            if st.button("Apply selected solutions", key="apply_solutions_v2"):
                selected = [s for s in solutions_catalog if s["name"] in sel_names]
                base_fill = st.session_state.get("fill_rate_pct", 60)
                agg = _aggregate_levers(selected, base_fill_rate=base_fill)
                # Push aggregated values into session so widgets auto-update
                st.session_state["aov_delta_pct"]        = int(agg["aov_delta_pct"])
                st.session_state["visit_freq_delta_pct"] = int(agg["visit_freq_delta_pct"])
                st.session_state["d_conv_pp"]            = int(agg["d_conv_pp"])
                st.session_state["attach_rate_delta_pct"]= int(agg["attach_rate_delta_pp"])
                st.session_state["avg_addon_value"]      = max(float(st.session_state.get("avg_addon_value", 120.0)),
                                                               float(agg["avg_addon_value"]))
                st.session_state["d_no_show_pp"]         = int(agg["d_no_show_pp"])
                st.session_state["csat_delta"]           = float(st.session_state.get("csat_delta", 0.0)) + float(agg["csat_delta"])
                st.session_state["d_handle_time_min"]    = int(agg["d_handle_time_min"])
                st.session_state["d_rework_pp"]          = int(agg["d_rework_pp"])
                st.session_state["fill_rate_pct"]        = int(agg["fill_rate_pct"])
                # If AOV direct % was provided, flip the radio below to “Direct”
                st.session_state["aov_mode_key"] = "Direct AOV increase (%)" if agg["aov_delta_pct"] else "Model via add-ons in Lever 4"
                st.success("Applied solutions → updated Core-8 levers. Fine-tune below.")
    else:
        st.caption("Tip: add `config/solutions.yml` to enable one-click solution presets.")

    # 1) Average Spend per Visit (AOV)
    with st.expander("1) Average Spend per Visit (AOV)", expanded=True):
        aov_mode = st.radio(
            "How should AOV be modeled?",
            ["Direct AOV increase (%)", "Model via add-ons in Lever 4"],
            index=1, key="aov_mode_key"
        )
        if aov_mode == "Direct AOV increase (%)":
            aov_delta_pct = st.slider(
                "Δ AOV (%)", 0, 25, st.session_state.get("aov_delta_pct", 5), 1, key="aov_delta_pct",
                help="Applies a % lift to Avg Ticket (AOV) for the scenario."
            )
        else:
            aov_delta_pct = 0  # handled via Lever 4
        st.caption("Evidence: add-on consistency, consult flow, segment spend patterns")

    # 2) Visit Frequency
    with st.expander("2) Visit Frequency", expanded=True):
        visit_freq_delta_pct = st.slider(
            "Δ visit frequency (scheduled visits %)", -10, 30,
            st.session_state.get("visit_freq_delta_pct", 0), 1, key="visit_freq_delta_pct",
            help="Approx. % change to scheduled visits/mo (pre-booking, plans, reminders)."
        )
        st.caption("Evidence: lapsed customer patterns, seasonality")

    # 3) New Customer Conversion
    with st.expander("3) New Customer Conversion", expanded=True):
        d_conv_pp = st.slider("Δ conversion (pp)", -20, 20,
                              st.session_state.get("d_conv_pp", 5), 1, key="d_conv_pp")
        st.caption("Evidence: speed-to-response, booking UX, consult structure")

    # 4) Upsell / Cross-sell
    with st.expander("4) Upsell / Cross-sell", expanded=True):
        if aov_mode == "Model via add-ons in Lever 4":
            attach_rate_delta_pct = st.slider(
                "Δ attach rate for add-ons (pp)", -20, 40,
                st.session_state.get("attach_rate_delta_pct", 5), 1, key="attach_rate_delta_pct"
            )
            avg_addon_value = st.number_input(
                "Avg add-on value (CAD)",
                value=float(st.session_state.get("avg_addon_value", 120.0)),
                min_value=0.0, step=10.0, key="avg_addon_value"
            )
            st.caption("Formula: Upsell Rev = baseline treated × Δ attach rate × add-on value")
        else:
            attach_rate_delta_pct = st.session_state.get("attach_rate_delta_pct", 0)
            avg_addon_value = float(st.session_state.get("avg_addon_value", 120.0))
            st.info("AOV is being modeled directly in Lever 1, so Upsell inputs are disabled here.")

    # 5) No-Show Reduction
    with st.expander("5) No-Show Reduction", expanded=True):
        d_no_show_pp = st.slider("Δ no-show (pp)", -20, 20,
                                 st.session_state.get("d_no_show_pp", -3), 1, key="d_no_show_pp")
        st.caption("Evidence: deposits, reminders, cancellation policy clarity")

    # 6) Retention / Repeat Rate
    with st.expander("6) Retention / Repeat Rate", expanded=True):
        st.caption("Modeled via CSAT → retention elasticity in Step 2.")
        _base  = float(st.session_state.get("base_csat", base_csat))
        _delta = float(st.session_state.get("csat_delta", csat_delta))
        _elas  = float(st.session_state.get("elasticity", elasticity))
        st.caption(f"Current: Baseline CSAT **{_base:.1f}**, CSAT delta **{_delta:+.2f}**, elasticity **{_elas:.2f}**.")

    # 7) Staff Productivity / Utilization
    with st.expander("7) Staff Productivity / Utilization", expanded=True):
        colp1, colp2, colp3 = st.columns(3)
        with colp1:
            d_handle_time_min = st.number_input(
                "Δ handle-time per visit (minutes)",
                value=st.session_state.get("d_handle_time_min", 3),
                min_value=-30, max_value=60, step=1, key="d_handle_time_min"
            )
        with colp2:
            d_rework_pp = st.slider("Δ rework rate (pp)", -30, 30,
                                    st.session_state.get("d_rework_pp", -4), 1, key="d_rework_pp")
        with colp3:
            fill_rate_pct = st.slider(
                "Capacity fill rate for recovered minutes (%)", 0, 100,
                st.session_state.get("fill_rate_pct", 60), 5, key="fill_rate_pct"
            )
        minutes_per_new_visit = st.number_input(
            "Minutes needed per new visit",
            value=int(st.session_state.get("minutes_per_new_visit", 35)),
            min_value=1, step=1
        )
        st.caption("Evidence: room flow, handoff delays, admin time. Minutes saved → extra capacity (at your fill rate).")

    # 8) Review Generation / Reputation (placeholder)
    with st.expander("8) Review Generation / Reputation", expanded=False):
        st.caption("Not yet wired to CAC/ROAS in this tab. Next: tie star rating → organic acquisition uplift / CAC reduction.")

    # ------------------ STEP 3.5: Costs ------------------
    st.markdown("### Costs")
    cc1, cc2 = st.columns(2)
    with cc1:
        tool_cost_mo = st.number_input("Tooling/platform cost (CAD/mo)", value=0, step=100, key="tool_cost_mo")
    with cc2:
        project_cost_one_time = st.number_input("One-time project cost (CAD)", value=0, step=500, key="project_cost_one_time")

    # ------------------ STEP 4: Calculations (baseline vs scenario) ------------------
    def _get(name, default):
        # read from session or fall back to local/default
        return st.session_state.get(name, default)

    # Core anchors
    visits_base           = int(_get("visits", visits))
    avg_ticket_base       = float(_get("avg_ticket", avg_ticket))
    avg_clv_used = float(_get("avg_clv", avg_clv_base))

    # Baseline rates
    conv_pct_val          = int(_get("conv_pct", consult_to_treatment_conv))
    no_show_pct_val       = int(_get("no_show_pct", no_show_rate))
    minutes_per_visit_val = int(_get("minutes_per_visit", minutes_per_visit))
    rework_pct_val        = int(_get("rework_pct", avg_rework_rate))
    avg_rework_minutes_val= int(_get("avg_rework_minutes", avg_rework_minutes))

    # Productivity / capacity
    fill_rate_pct_val     = int(_get("fill_rate_pct", fill_rate_pct))
    minutes_per_new_visit_val = int(_get("minutes_per_new_visit", minutes_per_new_visit))

    # Retention levers (Step 2)
    use_lifetime_val  = bool(_get("use_lifetime", use_lifetime))
    csat_delta_val    = float(_get("csat_delta", csat_delta))
    elasticity_val    = float(_get("elasticity", elasticity))

    # Upsell levers
    attach_rate_delta_pct_val = int(_get("attach_rate_delta_pct", attach_rate_delta_pct))
    avg_addon_value_val       = float(_get("avg_addon_value", avg_addon_value))

    # Costs
    tool_cost_mo_val          = float(_get("tool_cost_mo", tool_cost_mo))
    project_cost_one_time_val = float(_get("project_cost_one_time", project_cost_one_time))

    # Scenario deltas (Core-8)
    d_conv_pp_val          = int(_get("d_conv_pp", d_conv_pp))
    d_no_show_pp_val       = int(_get("d_no_show_pp", d_no_show_pp))
    d_handle_time_min_val  = int(_get("d_handle_time_min", d_handle_time_min))
    d_rework_pp_val        = int(_get("d_rework_pp", d_rework_pp))
    aov_mode_val           = _get("aov_mode_key", aov_mode)
    aov_delta_pct_val      = int(_get("aov_delta_pct", aov_delta_pct))
    visit_freq_delta_pct_val = int(_get("visit_freq_delta_pct", visit_freq_delta_pct))

    # If frequency lever used, bump scheduled visits for scenario only
    visits_scenario = int(round(visits_base * (1 + visit_freq_delta_pct_val/100.0)))

    # If direct AOV mode chosen, bump AOV for scenario only
    avg_ticket_scenario = avg_ticket_base * (1 + aov_delta_pct_val/100.0) if aov_mode_val == "Direct AOV increase (%)" else avg_ticket_base

    # ---- ROI calculator (single-source-of-truth) ----
    def _calc_roi(visits, avg_ticket, avg_clv,
                  conv_pct, no_show_pct,
                  minutes_per_visit, rework_pct, avg_rework_minutes,
                  d_conv_pp=0, d_no_show_pp=0, d_handle_time_min=0, d_rework_pp=0,
                  fill_rate_pct=60, minutes_per_new_visit=35,
                  csat_delta=0.0, elasticity=0.12, use_lifetime=True,
                  attach_rate_delta_pct=0, avg_addon_value=0.0,
                  tool_cost_mo=0.0, project_cost_one_time=0.0):

        conv = conv_pct/100.0
        no_show = no_show_pct/100.0
        rework_rate = rework_pct/100.0
        fill_rate = fill_rate_pct/100.0

        # Base volumes
        completed_visits = visits * (1 - no_show)
        treated = completed_visits * conv

        # A) no-show improvement
        improved_no_show = max(0.0, no_show + (d_no_show_pp/100.0))
        improved_completed = visits * (1 - improved_no_show)
        delta_completed_from_noshow = max(0.0, improved_completed - completed_visits)
        addl_treated_from_noshow = delta_completed_from_noshow * max(0.0, (conv + d_conv_pp/100.0))

        # B) conversion improvement on baseline
        addl_treated_from_conv = completed_visits * max(0.0, (conv + d_conv_pp/100.0) - conv)

        # C) minutes saved -> extra capacity
        minutes_saved_handle = completed_visits * max(0.0, d_handle_time_min)
        new_rework_rate = max(0.0, rework_rate + d_rework_pp/100.0)
        minutes_saved_rework = completed_visits * max(0.0, (rework_rate - new_rework_rate)) * avg_rework_minutes
        total_minutes_saved = minutes_saved_handle + minutes_saved_rework
        visit_capacity_from_minutes = (total_minutes_saved * fill_rate) / max(1, minutes_per_new_visit)
        addl_treated_from_capacity = visit_capacity_from_minutes * max(0.0, (conv + d_conv_pp/100.0))

        # Totals
        addl_treated_total = addl_treated_from_noshow + addl_treated_from_conv + addl_treated_from_capacity

        # Revenue components
        rev_lift_immediate = addl_treated_total * avg_ticket

        # Upsell / cross-sell (on baseline treated)
        attach_rate_delta = max(0.0, attach_rate_delta_pct/100.0)
        rev_lift_upsell = treated * attach_rate_delta * avg_addon_value

        # Lifetime effect
        uplift_factor = (max(0.0, csat_delta) * elasticity) if use_lifetime else 0.0
        rev_lift_lifetime = completed_visits * uplift_factor * float(avg_clv) * 0.40

        monthly_lift = rev_lift_immediate + rev_lift_upsell + rev_lift_lifetime
        monthly_cost = float(tool_cost_mo)
        net_monthly = monthly_lift - monthly_cost
        payback_months = (project_cost_one_time / net_monthly) if (project_cost_one_time > 0 and net_monthly > 0) else None

        return dict(
            completed_visits=completed_visits,
            treated=treated,
            addl_treated_from_noshow=addl_treated_from_noshow,
            addl_treated_from_conv=addl_treated_from_conv,
            addl_treated_from_capacity=addl_treated_from_capacity,
            addl_treated_total=addl_treated_total,
            rev_lift_immediate=rev_lift_immediate,
            rev_lift_upsell=rev_lift_upsell,
            rev_lift_lifetime=rev_lift_lifetime,
            monthly_lift=monthly_lift,
            monthly_cost=monthly_cost,
            net_monthly=net_monthly,
            total_minutes_saved=total_minutes_saved,
            fte_equiv_recovered=(total_minutes_saved/60.0)/160.0,
            payback_months=payback_months,
        )

    # ---- Build BASELINE (all deltas = 0) ----
    baseline = _calc_roi(
        visits=visits_base, avg_ticket=avg_ticket_base, avg_clv=avg_clv_base,
        conv_pct=conv_pct_val, no_show_pct=no_show_pct_val,
        minutes_per_visit=minutes_per_visit_val, rework_pct=rework_pct_val, avg_rework_minutes=avg_rework_minutes_val,
        d_conv_pp=0, d_no_show_pp=0, d_handle_time_min=0, d_rework_pp=0,
        fill_rate_pct=fill_rate_pct_val, minutes_per_new_visit=minutes_per_new_visit_val,
        csat_delta=0.0, elasticity=elasticity_val, use_lifetime=use_lifetime_val,
        attach_rate_delta_pct=0, avg_addon_value=avg_addon_value_val,
        tool_cost_mo=tool_cost_mo_val, project_cost_one_time=project_cost_one_time_val
    )

    # ---- Build SCENARIO (with your levers + Core-8) ----
    scenario = _calc_roi(
        visits=visits_scenario, avg_ticket=avg_ticket_scenario, avg_clv=avg_clv_base,
        conv_pct=conv_pct_val, no_show_pct=no_show_pct_val,
        minutes_per_visit=minutes_per_visit_val, rework_pct=rework_pct_val, avg_rework_minutes=avg_rework_minutes_val,
        d_conv_pp=d_conv_pp_val, d_no_show_pp=d_no_show_pp_val, d_handle_time_min=d_handle_time_min_val, d_rework_pp=d_rework_pp_val,
        fill_rate_pct=fill_rate_pct_val, minutes_per_new_visit=minutes_per_new_visit_val,
        csat_delta=csat_delta_val, elasticity=elasticity_val, use_lifetime=use_lifetime_val,
        attach_rate_delta_pct=attach_rate_delta_pct_val, avg_addon_value=avg_addon_value_val,
        tool_cost_mo=tool_cost_mo_val, project_cost_one_time=project_cost_one_time_val
    )

    # --- Monthly revenue baseline & % lift (add right after baseline/scenario) ---
    with st.expander("Baseline revenue reference (for % lift)", expanded=True):
        annual_rev_baseline_val = st.number_input(
            "Baseline annual revenue (CAD)",
            value=1_500_000.0, step=50_000.0, key="annual_rev_baseline_val"
        )

    baseline_monthly_rev = annual_rev_baseline_val / 12.0
    new_monthly_rev = baseline_monthly_rev + float(scenario["monthly_lift"])
    pct_lift_vs_baseline = (float(scenario["monthly_lift"]) / baseline_monthly_rev) * 100.0 if baseline_monthly_rev > 0 else 0.0

    c_rev1, c_rev2, c_rev3 = st.columns(3)
    with c_rev1:
        st.metric("Baseline monthly revenue", f"${baseline_monthly_rev:,.0f}")
    with c_rev2:
        st.metric("New monthly revenue (modeled)", f"${new_monthly_rev:,.0f}",
                  delta=f"+${scenario['monthly_lift']:,.0f}")
    with c_rev3:
        st.metric("% lift vs baseline", f"{pct_lift_vs_baseline:.1f}%")


    # ---- Comparison table ----
    import pandas as pd

    def _money(x):
        try: return f"${x:,.0f}"
        except: return x

    def _num(x):
        try: return f"{x:,.0f}"
        except: return x

    rows = [
        ("Scheduled visits / mo (input)", visits_base,           visits_scenario,        False),
        ("Avg ticket / AOV (CAD)",        avg_ticket_base,       avg_ticket_scenario,    True),

        ("Completed visits / mo",         baseline["completed_visits"], scenario["completed_visits"], False),
        ("Treated visits / mo",           baseline["treated"],          scenario["treated"],          False),
        ("Added treated / mo",            0,                              scenario["addl_treated_total"], False),

        ("Immediate lift (CAD/mo)",       0,                              scenario["rev_lift_immediate"], True),
        ("Upsell lift (CAD/mo)",          0,                              scenario["rev_lift_upsell"],   True),
        ("Lifetime lift (CAD/mo)",        0,                              scenario["rev_lift_lifetime"], True),
        ("Monthly lift (CAD/mo)",         0,                              scenario["monthly_lift"],      True),

        ("Monthly cost (CAD/mo)",         baseline["monthly_cost"],      scenario["monthly_cost"],      True),
        ("Net monthly (CAD/mo)",          baseline["net_monthly"],       scenario["net_monthly"],       True),
    ]
    cmp_df = pd.DataFrame(rows, columns=["Metric","Baseline","With solutions","is_money"])
    cmp_df["Δ"] = cmp_df["With solutions"] - cmp_df["Baseline"]

    pretty = cmp_df.copy()
    pretty["Baseline"]       = pretty.apply(lambda r: _money(r["Baseline"]) if r["is_money"] else _num(r["Baseline"]), axis=1)
    pretty["With solutions"] = pretty.apply(lambda r: _money(r["With solutions"]) if r["is_money"] else _num(r["With solutions"]), axis=1)
    pretty["Δ"]              = pretty.apply(lambda r: _money(r["Δ"]) if r["is_money"] else _num(r["Δ"]), axis=1)

    st.markdown("### Baseline vs. With solutions")
    st.dataframe(pretty[["Metric","Baseline","With solutions","Δ"]], use_container_width=True)

    # ---- Headline KPIs (scenario) ----
    st.markdown("### Results (with solutions)")
    k1, k2, k3, k4 = st.columns(4)
    with k1: st.metric("Added treated / mo", f"{int(scenario['addl_treated_total']):,}")
    with k2: st.metric("Immediate lift (CAD/mo)", f"{scenario['rev_lift_immediate']:,.0f}")
    with k3: st.metric("Lifetime lift (CAD/mo)", f"{scenario['rev_lift_lifetime']:,.0f}")
    with k4: st.metric("Net monthly (CAD)", f"{scenario['net_monthly']:,.0f}")

    h1, h2, h3 = st.columns(3)
    with h1: st.metric("Minutes saved / mo", f"{int(scenario['total_minutes_saved']):,}")
    with h2: st.metric("FTE equiv. recovered", f"{scenario['fte_equiv_recovered']:.2f}")
    with h3: st.metric("Payback (months)", f"{scenario['payback_months']:.1f}" if scenario['payback_months'] else "—")

# =========================================================
# Tab 5 (new): Market Expansion — TAM & Share Uplift (B2B2C)
# =========================================================
with tab_market:
    try:
        st.subheader("Market Expansion — Wearable TAM & Share Uplift (B2B2C)")

        # --- small, local helpers (avoids NameError if not defined elsewhere) ---
        import os, json
        try:
            import yaml  # PyYAML
        except Exception:
            yaml = None

        def load_markets_yaml(path: str = "config/markets.yml") -> dict:
            if not os.path.exists(path) or yaml is None:
                return {}
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f) or {}
                return data.get("markets", {}) or {}
            except Exception:
                return {}

        def load_market_levers_yaml(path: str = "config/market_levers.yml") -> dict:
            if not os.path.exists(path) or yaml is None:
                return {}
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f) or {}
                return data.get("levers", {}) or {}
            except Exception:
                return {}

        def load_market_scenarios_yaml(path: str = "config/market_scenarios.yml") -> dict:
            if not os.path.exists(path) or yaml is None:
                return {}
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = yaml.safe_load(f) or {}
                return data.get("scenarios", {}) or {}
            except Exception:
                return {}

        def _intensity_key(region: str, lever_key: str) -> str:
            return f"mx_intensity::{region}::{lever_key}"

        def _lever_effect_pp(levers_cfg: dict, lever_key: str, mode: str, intensity_0_1: float) -> float:
            rng = (levers_cfg.get(lever_key, {}).get("effect_on_share_pp") or [0.0, 0.0])
            lo = float(rng[0]) if len(rng) > 0 else 0.0
            hi = float(rng[1]) if len(rng) > 1 else lo
            if mode == "Conservative":
                base = lo
            elif mode == "Aggressive":
                base = hi
            else:
                base = (lo + hi) / 2.0
            intensity = max(0.0, min(1.0, float(intensity_0_1)))
            return base * intensity  # pp

        # --- load configs ---
        markets_cfg   = load_markets_yaml("config/markets.yml")
        levers_cfg    = load_market_levers_yaml("config/market_levers.yml")
        scenarios_cfg = load_market_scenarios_yaml("config/market_scenarios.yml")

        # Quick debug row so you know what loaded on Cloud
        st.caption(f"Loaded: markets={len(markets_cfg)} | levers={len(levers_cfg)} | scenarios={len(scenarios_cfg)}")

        if not markets_cfg:
            st.info("No markets config found. Add **config/markets.yml** (see example in the spec).")
            st.stop()

        # --- 1) Global scenario settings ---
        c1, c2, c3 = st.columns([1,1,2])
        with c1:
            s1 = st.number_input("Scenario 1: target share (%)", value=1, min_value=0, max_value=50, step=1, key="mx_s1")
        with c2:
            s2 = st.number_input("Scenario 2: target share (%)", value=3, min_value=0, max_value=50, step=1, key="mx_s2")
        with c3:
            mode = st.radio("Assumption mode (lever ranges)", ["Conservative","Typical","Aggressive"], horizontal=True, key="mx_mode")

        targets = sorted({int(s1), int(s2)})

        # --- 2) Regions & (optional) scenario preset ---
        region_keys = list(markets_cfg.keys())
        chosen_regions = st.multiselect("Regions to model", options=region_keys, default=region_keys, key="mx_regions_sel")

        scenames = ["(None)"] + (list(scenarios_cfg.keys()) if scenarios_cfg else [])
        sel_scn = st.selectbox("Load saved lever intensities (optional)", scenames, index=0, key="mx_scenario_sel")

        if sel_scn != "(None)" and scenarios_cfg:
            picked = scenarios_cfg.get(sel_scn, {}).get("regions", {})
            for r_key, r_def in picked.items():
                for lever_key, val in (r_def.get("lever_intensity") or {}).items():
                    st.session_state[_intensity_key(r_key, lever_key)] = float(val)

        # Early empty-state guard
        if not chosen_regions:
            st.warning("Select at least one region to run the model.")
            st.stop()

        # --- 3) Levers UI ---
        if not levers_cfg:
            st.caption("Tip: add **config/market_levers.yml** to enable lever modeling.")
        else:
            st.markdown("### Levers (per region)")
            for r in chosen_regions:
                st.markdown(f"**{markets_cfg[r].get('label', r)}**")
                cols = st.columns(3)
                i = 0
                for lever_key, meta in levers_cfg.items():
                    label = meta.get("label", lever_key)
                    help_txt = f"Effect @100%: {meta.get('effect_on_share_pp', [0,0])} pp"
                    k = _intensity_key(r, lever_key)
                    tmp_key = k + "::pct"
                    cur_pct = int(round(100 * float(st.session_state.get(k, 0.0))))
                    with cols[i % 3]:
                        pct_val = st.slider(label, 0, 100, cur_pct, 5, help=help_txt, key=tmp_key)
                        st.session_state[k] = pct_val / 100.0
                    i += 1

        st.divider()

        # --- 4) Compute ---
        import pandas as pd
        rows_proj, rows_targets = [], []

        for r in chosen_regions:
            m = markets_cfg[r]
            label = m.get("label", r)
            currency = m.get("currency", "USD")
            tam = float(m.get("wearable_tam", 0.0))
            reachable_pct = float(m.get("reachable_pct", 1.0))
            baseline_share = float(m.get("baseline_share_pct", 0.0))
            gm_pct = float(m.get("gross_margin_pct", 0.0))

            reachable_rev = tam * reachable_pct
            base_rev = reachable_rev * (baseline_share / 100.0)
            base_gp = base_rev * gm_pct

            uplift_pp = 0.0
            for lever_key in levers_cfg.keys():
                intensity = float(st.session_state.get(_intensity_key(r, lever_key), 0.0))
                uplift_pp += _lever_effect_pp(levers_cfg, lever_key, mode, intensity)

            projected_share = max(0.0, baseline_share + uplift_pp)
            proj_rev = reachable_rev * (projected_share / 100.0)
            proj_gp = proj_rev * gm_pct

            rows_proj.append({
                "Region": label,
                "Currency": currency,
                "Baseline_share_pct": baseline_share,
                "Projected_share_pct": projected_share,
                "Reachable_TAM": reachable_rev,
                "Baseline_revenue": base_rev,
                "Projected_revenue": proj_rev,
                "Incremental_revenue": proj_rev - base_rev,
                "Baseline_gross_profit": base_gp,
                "Projected_gross_profit": proj_gp,
                "Incremental_gross_profit": proj_gp - base_gp,
                "Assumption_mode": mode,
            })

            for t in targets:
                t_share = float(t)
                t_rev = reachable_rev * (t_share / 100.0)
                t_gp  = t_rev * gm_pct
                rows_targets.append({
                    "Region": label,
                    "Currency": currency,
                    "Baseline_share_pct": baseline_share,
                    "Target_share_pct": t_share,
                    "Reachable_TAM": reachable_rev,
                    "Baseline_revenue": base_rev,
                    "Target_revenue": t_rev,
                    "Incremental_revenue": t_rev - base_rev,
                    "Incremental_gross_profit": t_gp - base_gp,
                })

        # --- 5) Render ---
        def _money(x):
            try: return f"{x:,.0f}"
            except: return x

        proj_df = pd.DataFrame(rows_proj)
        targ_df = pd.DataFrame(rows_targets)

        if proj_df.empty and targ_df.empty:
            st.info("No rows to display yet. Check your regions and lever selections.")
            st.stop()

        if not proj_df.empty:
            view = proj_df.copy()
            for c in ["Reachable_TAM","Baseline_revenue","Projected_revenue",
                      "Incremental_revenue","Baseline_gross_profit",
                      "Projected_gross_profit","Incremental_gross_profit"]:
                view[c] = view[c].apply(_money)
            st.markdown("### Lever-driven Projection")
            st.dataframe(
                view[["Region","Currency","Assumption_mode",
                      "Baseline_share_pct","Projected_share_pct",
                      "Baseline_revenue","Projected_revenue",
                      "Incremental_revenue","Incremental_gross_profit"]],
                use_container_width=True
            )

        if not targ_df.empty:
            view2 = targ_df.copy()
            for c in ["Reachable_TAM","Baseline_revenue","Target_revenue",
                      "Incremental_revenue","Incremental_gross_profit"]:
                view2[c] = view2[c].apply(_money)
            st.markdown("### TAM & Share Scenarios (independent of levers)")
            st.dataframe(
                view2[["Region","Currency","Baseline_share_pct","Target_share_pct",]()]()

